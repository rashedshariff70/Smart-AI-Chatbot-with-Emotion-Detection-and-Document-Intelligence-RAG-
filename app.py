from flask import Flask, render_template, request, jsonify, session
import os
import json
import re
from werkzeug.utils import secure_filename
import tempfile

# PDF/DOC reading
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB

UPLOAD_FOLDER = tempfile.mkdtemp()
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

doc_store: dict[str, str] = {}
conv_store: dict[str, list] = {}

# ─────────────────────────────────────────
# 🔧 LLM PLACEHOLDER (SAFE FOR GITHUB)
# ─────────────────────────────────────────
def call_llm(system_prompt: str, user_message: str) -> str:
    """
    Replace this function with your own LLM integration.
    Example:
        - OpenAI
        - Groq
        - Ollama (local)
        - HuggingFace
    """
    return "LLM response placeholder. Integrate your model here."


# ─────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(filepath: str, ext: str) -> str:
    if ext == "txt":
        with open(filepath, "r", errors="ignore") as f:
            return f.read()

    if ext == "pdf":
        if not PDF_SUPPORT:
            raise RuntimeError("Install PyPDF2")
        text = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)

    if ext == "docx":
        if not DOCX_SUPPORT:
            raise RuntimeError("Install python-docx")
        doc = DocxDocument(filepath)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    return ""


# ─────────────────────────────────────────
# Routes – Pages
# ─────────────────────────────────────────

@app.route("/")
def home():
    return render_template("home.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/prediction")
def prediction():
    return render_template("prediction.html")


# ─────────────────────────────────────────
# API – Stage 1 : Emotion Detection
# ─────────────────────────────────────────

@app.route("/api/emotion", methods=["POST"])
def api_emotion():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()

    if not text:
        return jsonify({"error": "No text provided"}), 400

    # Placeholder response
    result = {
        "primary_emotion": "neutral",
        "confidence": 75,
        "emotions": {
            "joy": 10,
            "sadness": 5,
            "anger": 5,
            "fear": 5,
            "surprise": 5,
            "disgust": 5,
            "neutral": 65
        },
        "sentiment": "neutral",
        "intensity": "low",
        "summary": "LLM integration required for real emotion analysis."
    }

    return jsonify(result)


# ─────────────────────────────────────────
# API – Stage 2 : RAG Document Chatbot
# ─────────────────────────────────────────

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]

    if not file or file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[1].lower()
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        text = extract_text(filepath, ext)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    if not text.strip():
        return jsonify({"error": "No text extracted"}), 400

    sid = session.get("sid") or os.urandom(8).hex()
    session["sid"] = sid

    doc_store[sid] = text[:15000]

    return jsonify({
        "success": True,
        "chars": len(text),
        "preview": text[:300]
    })


@app.route("/api/rag", methods=["POST"])
def api_rag():
    data = request.get_json(silent=True) or {}
    question = (data.get("question") or "").strip()
    sid = session.get("sid", "")

    if not question:
        return jsonify({"error": "No question"}), 400

    doc_text = doc_store.get(sid, "")

    if not doc_text:
        return jsonify({"error": "Upload document first"}), 400

    # Placeholder answer
    answer = "This is a placeholder answer. Connect your LLM to enable document-based Q&A."

    return jsonify({"answer": answer})


# ─────────────────────────────────────────
# API – Stage 3 : Chatbot
# ─────────────────────────────────────────

@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()

    if not message:
        return jsonify({"error": "No message"}), 400

    sid = session.get("sid") or os.urandom(8).hex()
    session["sid"] = sid

    history = conv_store.setdefault(sid, [])
    history.append({"role": "user", "content": message})

    # Placeholder chatbot reply
    reply = "Chatbot placeholder. Integrate your own LLM model here."

    history.append({"role": "assistant", "content": reply})

    return jsonify({"reply": reply})


# ─────────────────────────────────────────

if __name__ == "__main__":
    app.run(debug=True, port=5000)